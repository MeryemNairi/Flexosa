from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import StreamingResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
import base64
from datetime import datetime
import os
from typing import Dict
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from langchain_groq import ChatGroq
from langchain.schema import HumanMessage, SystemMessage
from langchain_community.chat_models import ChatLiteLLM
from langchain.prompts import ChatPromptTemplate
from dotenv import load_dotenv
from fastapi import Request

# Load environment variables
load_dotenv()

# Initialize FastAPI app
app = FastAPI(title="X-Ray Report Generator")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Templates
templates = Jinja2Templates(directory="templates")

# Prompt templates
VISION_ANALYSIS_TEMPLATE = """
You are an AI assistant tasked with analyzing an X-ray image for diagnostic purposes. 
Analyze the image and provide detailed information about:

1. Clinical Findings:
   - Describe any visible abnormalities
   - Note bone structure and alignment
   - Identify any soft tissue changes

2. Technical Assessment:
   - Image quality
   - Patient positioning
   - Exposure factors

3. Anatomical Structures:
   - Identify and describe key anatomical structures
   - Note any variations from normal anatomy

4. Pathological Findings:
   - Describe any pathological changes
   - Note any degenerative changes
   - Identify any acute findings

Provide your analysis in a structured format that can be easily processed by another AI model.
"""

REPORT_GENERATION_TEMPLATE = """
You are a medical report writer. Based on the provided X-ray analysis, create a comprehensive medical report.
The report should be professional, clear, and well-structured.

Patient Information:
Name: {patient_name}
Date of Birth: {patient_dob}
Date of Examination: {exam_date}

Previous Analysis:
{analysis}

Please generate a formal medical report that includes:

1. Clinical History
2. Technique
3. Findings
4. Impression
5. Recommendations

Use medical terminology appropriately but ensure the report remains comprehensible.
"""

class LLMService:
    def _init_(self):
        self.api_key = "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"

    def encode_image(self, image_content: bytes) -> str:
        return base64.b64encode(image_content).decode('utf-8')

    async def analyze_image(self, image_content: bytes) -> str:
        from groq import Groq
        
        client = Groq(api_key=self.api_key)
        base64_image = self.encode_image(image_content)
        
        completion = client.chat.completions.create(
            model="llama-3.2-90b-vision-preview",
            messages=[
                
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": VISION_ANALYSIS_TEMPLATE
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
            temperature=0.7,
            max_tokens=1024,
            top_p=1,
            stream=False
        )
        print(completion.choices[0].message.content)
        return completion.choices[0].message.content

    async def generate_report(self, analysis: str, patient_info: Dict[str, str]) -> str:
        from groq import Groq
        
        client = Groq(api_key=self.api_key)
        
        prompt = REPORT_GENERATION_TEMPLATE.format(
            patient_name=patient_info["name"],
            patient_dob=patient_info["dob"],
            exam_date=datetime.now().strftime("%Y-%m-%d"),
            analysis=analysis
        )
        
        completion = client.chat.completions.create(
            model="mixtral-8x7b-32768",  # Using Mixtral for text generation
            messages=[
                {
                    "role": "system",
                    "content": "You are a medical report writer. Generate a professional and comprehensive medical report."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.3,
            max_tokens=2048,
            top_p=1,
            stream=False
        )
        
        return completion.choices[0].message.content

class DocumentService:
    @staticmethod
    def create_report_document(content: str, patient_info: dict) -> BytesIO:
        doc = Document()
        
        # Add header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "DHA Diagnostic Center"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add title
        title = doc.add_heading('X-Ray Report', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add patient information
        doc.add_paragraph(f"Patient Name: {patient_info['name']}")
        doc.add_paragraph(f"Date of Birth: {patient_info['dob']}")
        doc.add_paragraph(f"Date of Examination: {datetime.now().strftime('%Y-%m-%d')}")
        
        # Add horizontal line
        doc.add_paragraph('_' * 50)
        
        # Add content sections
        sections = content.split('\n\n')
        for section in sections:
            if section.strip():
                if section.upper().startswith(('FINDINGS:', 'IMPRESSION:', 'TECHNIQUE:')):
                    # Add section headers in bold
                    p = doc.add_paragraph()
                    p.add_run(section.split(':')[0] + ':').bold = True
                    p.add_run(section.split(':', 1)[1])
                else:
                    doc.add_paragraph(section)
        
        # Add footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to BytesIO
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer

# Initialize services
llm_service = LLMService()
doc_service = DocumentService()

@app.get("/", response_class=HTMLResponse)
async def get_upload_page(request: Request):
    return templates.TemplateResponse(
        "upload.html",
        {"request": request}
    )

@app.post("/api/process-image")
async def process_image(
    image: UploadFile = File(...),
    patient_name: str = Form(...),
    patient_dob: str = Form(...)
):
    try:
        # Read image content
        image_content = await image.read()
        
        # Step 1: Analyze image with Llama Vision
        print("Analyzing image...")
        analysis = await llm_service.analyze_image(image_content)
        print("Analysis complete")
        
        # Step 2: Generate report with Groq
        print("Generating report...")
        patient_info = {
            "name": patient_name,
            "dob": patient_dob
        }
        report_content = await llm_service.generate_report(analysis, patient_info)
        print("Report generated")
        
        # Step 3: Generate document
        print("Creating document...")
        doc_buffer = doc_service.create_report_document(
            report_content,
            patient_info
        )
        print("Document created")
        
        # Return the document
        return StreamingResponse(
            doc_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": "attachment; filename=xray_report.docx"
            }
        )
        
    except Exception as e:
        print(f"Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
